import os
import re
import PyPDF2
import uuid
from langchain.text_splitter import RecursiveCharacterTextSplitter
from langchain.docstore.document import Document
import streamlit as st
from config import PARENT_CHUNK_SIZE, PARENT_CHUNK_OVERLAP, CHILD_CHUNK_SIZE, CHILD_CHUNK_OVERLAP
import docx
import concurrent.futures
from concurrent.futures import ThreadPoolExecutor
import time

def process_text_file(uploaded_file):
    """Xử lý file text"""
    try:
        text = str(uploaded_file.read(), "utf-8")
        return Document(page_content=text, metadata={"source": uploaded_file.name})
    except Exception as e:
        st.error(f"Lỗi khi đọc file text '{uploaded_file.name}': {e}")
        return None

def process_pdf_file(uploaded_file):
    """Xử lý file PDF với PyPDF2"""
    try:
        pdf_reader = PyPDF2.PdfReader(uploaded_file)
        text = ""
        for page in pdf_reader.pages:
            text += page.extract_text() or ""
        return Document(page_content=text, metadata={"source": uploaded_file.name})
    except Exception as e:
        st.error(f"Lỗi khi đọc file PDF '{uploaded_file.name}': {e}")
        return None

def process_docx_file(uploaded_file):
    """Xử lý file DOCX"""
    try:
        doc = docx.Document(uploaded_file)
        text = "\n\n".join([paragraph.text for paragraph in doc.paragraphs])
        return Document(page_content=text, metadata={"source": uploaded_file.name})
    except Exception as e:
        st.error(f"Lỗi khi đọc file DOCX '{uploaded_file.name}': {e}")
        return None

def process_single_file(uploaded_file):
    """Xử lý một file tải lên và trả về Document"""
    if uploaded_file is not None:
        file_name = uploaded_file.name
        file_extension = os.path.splitext(file_name)[1].lower()
        
        # Chọn hàm xử lý phù hợp theo loại file
        if file_extension == ".txt":
            doc = process_text_file(uploaded_file)
        elif file_extension == ".pdf":
            doc = process_pdf_file(uploaded_file)
        elif file_extension in [".doc", ".docx"]:
            doc = process_docx_file(uploaded_file)
        else:
            st.error(f"Định dạng file '{file_name}' không được hỗ trợ.")
            return None
        
        # Kiểm tra và xử lý văn bản
        if doc and doc.page_content.strip():
            # Làm sạch văn bản
            doc.page_content = re.sub(r'\s+', ' ', doc.page_content).strip()
            return doc
        else:
            if doc: # Nếu tạo được doc nhưng không có nội dung
                st.warning(f"File '{file_name}' không có nội dung.")
            return None
    return None

def process_uploaded_files(uploaded_files):
    """
    Đọc các file được tải lên, làm sạch, và chia thành các parent và child chunks.
    Trả về một tuple: (parent_chunks, child_chunks)
    """
    print("[document_processor] Bắt đầu xử lý và chia chunk parent/child...")
    start_time = time.time()
    
    if not isinstance(uploaded_files, list):
        uploaded_files = [uploaded_files]

    if not uploaded_files:
        print("[document_processor] Không có file nào được cung cấp.")
        return None, None

    raw_docs_with_source = []
    file_errors = []
    
    # Xử lý song song các file với ThreadPoolExecutor
    with ThreadPoolExecutor(max_workers=min(10, len(uploaded_files))) as executor:
        # Map xử lý từng file
        future_to_file = {executor.submit(process_single_file, file): file for file in uploaded_files}
        for future in concurrent.futures.as_completed(future_to_file):
            file = future_to_file[future]
            try:
                doc = future.result()
                if doc:
                    raw_docs_with_source.append(doc)
                    print(f"[document_processor] Đã xử lý thành công file: {file.name} ({len(doc.page_content)} ký tự)")
                else:
                    file_errors.append(f"Không thể xử lý nội dung file: {file.name}")
            except Exception as e:
                print(f"[document_processor] Lỗi khi xử lý file '{file.name}': {e}")
                file_errors.append(f"Lỗi xử lý file {file.name}: {str(e)}")
    
    if file_errors:
        error_message = "\n".join(file_errors)
        st.error(f"Có lỗi khi xử lý một số file:\n{error_message}")
    
    if not raw_docs_with_source:
        print("[document_processor] Không có văn bản nào được trích xuất.")
        return None, None
        
    # --- LOGIC CỦA PARENT DOCUMENT ---
    try:
        # 1. Chia thành các parent chunks
        parent_splitter = RecursiveCharacterTextSplitter(
            chunk_size=PARENT_CHUNK_SIZE, 
            chunk_overlap=PARENT_CHUNK_OVERLAP
        )
        parent_chunks = parent_splitter.split_documents(raw_docs_with_source)
        print(f"[document_processor] Đã chia thành {len(parent_chunks)} parent chunks.")
    
        # 2. Gán ID duy nhất cho mỗi parent chunk để liên kết
        id_key = "parent_id"
        for i, p_chunk in enumerate(parent_chunks):
            p_chunk.metadata[id_key] = str(uuid.uuid4())
            p_chunk.metadata["chunk_id"] = i  # Thêm chunk_id để giữ khả năng tương thích với code cũ
    
        # 3. Chia parent chunks thành các child chunks
        child_splitter = RecursiveCharacterTextSplitter(
            chunk_size=CHILD_CHUNK_SIZE, 
            chunk_overlap=CHILD_CHUNK_OVERLAP
        )
        
        child_chunks = []
        for p_chunk in parent_chunks:
            # Tách nội dung của parent chunk
            _child_docs_content = child_splitter.split_text(p_chunk.page_content)
            
            # Gán metadata và tạo Document cho child chunk
            for j, _child_content in enumerate(_child_docs_content):
                child_metadata = p_chunk.metadata.copy()
                child_metadata["child_chunk_id"] = j  # ID riêng cho child để phân biệt
                child_doc = Document(page_content=_child_content, metadata=child_metadata)
                child_chunks.append(child_doc)
    
            print(f"[document_processor] Đã chia thành {len(child_chunks)} child chunks.")
            print(f"[document_processor] Xử lý tài liệu hoàn tất trong {time.time() - start_time:.2f} giây.")
            
            # Hàm giờ trả về parent và child chunks
            return parent_chunks, child_chunks
    except Exception as e:
        print(f"[document_processor] Lỗi nghiêm trọng khi chia chunks: {e}")
        st.error(f"Lỗi khi xử lý tài liệu: {str(e)}")
        return None, None

# Hàm mới - xử lý adaptive dựa trên kích thước file
def process_uploaded_files_adaptive(uploaded_files):
    """
    Xử lý tài liệu với chiến lược thích nghi dựa trên kích thước file.
    Với file nhỏ, sử dụng chunking đơn giản; với file lớn, dùng parent-child chunking.
    """
    # Tính tổng kích thước các file
    total_size = sum(file.size for file in uploaded_files) if isinstance(uploaded_files, list) else uploaded_files.size
    
    # Ngưỡng 100KB - với file nhỏ hơn, dùng chunking đơn giản
    print(f"[document_processor] Tổng kích thước file: {total_size/1024:.1f}KB")
    
    if total_size < 100 * 1024:  # Dưới 100KB
        print("[document_processor] File nhỏ, xử lý với chunking đơn giản...")
        return process_simple_chunking(uploaded_files)
    else:
        print("[document_processor] File lớn, xử lý với parent-child chunking đầy đủ...")
        return process_uploaded_files(uploaded_files)

# Hàm mới - chunking đơn giản cho file nhỏ
def process_simple_chunking(uploaded_files):
    """Xử lý chunking đơn giản cho file nhỏ, không chia parent-child."""
    start_time = time.time()
    
    if not isinstance(uploaded_files, list):
        uploaded_files = [uploaded_files]

    raw_docs = []
    for file in uploaded_files:
        doc = process_single_file(file)
        if doc:
            raw_docs.append(doc)
    
    if not raw_docs:
        return None, None
        
    # Tạo chunks đơn giản
    splitter = RecursiveCharacterTextSplitter(
        chunk_size=CHILD_CHUNK_SIZE,
        chunk_overlap=CHILD_CHUNK_OVERLAP
    )
    
    child_chunks = splitter.split_documents(raw_docs)
    
    # Gán metadata
    for i, chunk in enumerate(child_chunks):
        chunk.metadata["parent_id"] = chunk.metadata.get("source", "") + f"_chunk_{i}"
        chunk.metadata["chunk_id"] = i
    
    # Với file nhỏ, parent chunks là chính child chunks
    parent_chunks = child_chunks.copy()
    
    print(f"[document_processor] Xử lý đơn giản hoàn tất trong {time.time() - start_time:.2f} giây.")
    print(f"[document_processor] Tạo {len(child_chunks)} chunks.")
    
    return parent_chunks, child_chunks

# (Code for process_uploaded_files will go here) 